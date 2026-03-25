from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE_DIR = Path(__file__).resolve().parent
DOCS_DIR = BASE_DIR / "docs"
OUTPUT = DOCS_DIR / "informe_tecnico.docx"


def heading_center(doc: Document, text: str, level: int = 0) -> None:
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)

    doc = Document()
    heading_center(doc, "Informe técnico del caso práctico N.º 01", 0)
    heading_center(doc, "Business Corporation - Sistema de Recursos Humanos", 1)

    p = doc.add_paragraph()
    p.add_run("Fecha: ").bold = True
    p.add_run(str(date.today()))

    doc.add_heading("1. Resumen del proyecto", level=1)
    doc.add_paragraph(
        "Se desarrolló una solución modular en Python para gestionar la estructura de recursos humanos de Business Corporation. "
        "La interfaz fue implementada con Streamlit y la lógica del negocio se organizó mediante Programación Orientada a Objetos."
    )

    doc.add_heading("2. Requisitos cubiertos", level=1)
    for item in [
        "Uso de clases, objetos, herencia, encapsulamiento y polimorfismo.",
        "Almacenamiento de trabajadores en una lista de objetos.",
        "Creación de un gerente, cinco jefes de área, asistentes y técnicos.",
        "Métodos get_resumen(), get_jefe_inmediato() y get_estado().",
        "Getters y setters para facilitar mantenimiento.",
        "Interfaz lista para ejecución local y despliegue en Streamlit Cloud.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3. Explicación de archivos", level=1)
    archivos = {
        "app.py": "Archivo principal de la interfaz Streamlit.",
        "models/trabajador.py": "Clase abstracta base con atributos encapsulados.",
        "models/roles.py": "Subclases de trabajadores con polimorfismo.",
        "services/repository.py": "Gestión del repositorio en memoria y reglas del negocio.",
        "services/data_factory.py": "Carga inicial de datos del caso.",
        "services/reporting.py": "Conversión a DataFrames e indicadores para la vista.",
        "utils/helpers.py": "Funciones auxiliares para la interfaz y decisiones de diseño.",
        "tests/test_logic.py": "Pruebas automáticas del comportamiento principal.",
        "scripts/setup_env.py": "Creación del entorno virtual e instalación de dependencias.",
        "requirements.txt": "Dependencias mínimas de ejecución para Streamlit Cloud.",
        "requirements-dev.txt": "Dependencias de desarrollo y documentación.",
    }
    for archivo, descripcion in archivos.items():
        p = doc.add_paragraph()
        p.add_run(f"{archivo}: ").bold = True
        p.add_run(descripcion)

    doc.add_heading("4. Corrección aplicada para Streamlit", level=1)
    doc.add_paragraph(
        "Se dejó el archivo requirements.txt con dependencias mínimas de ejecución para evitar despliegues lentos o bloqueados en Streamlit Cloud. "
        "Las dependencias de desarrollo fueron separadas en requirements-dev.txt."
    )

    doc.add_heading("5. Entorno virtual y ejecución", level=1)
    for comando in [
        "python scripts/setup_env.py",
        ".venv/bin/streamlit run app.py",
        ".venv/bin/pytest -q",
        ".venv/bin/python generate_report.py",
    ]:
        doc.add_paragraph(comando, style="List Bullet")

    doc.add_heading("6. Observación", level=1)
    doc.add_paragraph(
        "El enunciado menciona cuatro áreas concretas, pero también solicita cinco jefes de área. "
        "Para mantener coherencia y cumplir el mínimo exigido se añadió el área de Finanzas."
    )

    doc.add_heading("7. Repositorio y despliegue", level=1)
    doc.add_paragraph(
        "El proyecto queda preparado para subir a GitHub y desplegarse en Streamlit Community Cloud indicando app.py como archivo principal."
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
